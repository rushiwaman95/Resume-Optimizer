# app.py — Resume Optimizer (Production Ready)
# pip install flask python-docx requests bleach reportlab

import io
import re
import json
import hashlib
import time
import threading
import traceback
from html.parser import HTMLParser

import bleach
import docx
import requests as req
from flask import Flask, jsonify, render_template_string, request, send_file
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    ListFlowable, ListItem, Paragraph,
    SimpleDocTemplate, Spacer, Flowable
)

# ═══════════════════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════════════════
RESUME_MAX  = 15_000
JD_MAX      = 8_000
RESUME_MIN  = 100
API_TIMEOUT = 120

ALLOWED_TAGS  = ["h1","h2","h3","p","strong","em","ul","li","br"]
ALLOWED_ATTRS: dict = {}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024

# ── CORS ─────────────────────────────────────────────────────
@app.after_request
def cors(r):
    r.headers["Access-Control-Allow-Origin"]  = "*"
    r.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    r.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return r

# ═══════════════════════════════════════════════════════════════
# HTML PAGE
# ═══════════════════════════════════════════════════════════════
PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>Resume Optimizer — ATS Ready</title>
<style>
/* ── reset ── */
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}

/* ── base ── */
body{
  font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,
              Helvetica,Arial,sans-serif;
  background:#f0f2f6;
  color:#1a1a2e;
  min-height:100vh;
  display:flex;
  font-size:14px;
}

/* ══════════ SIDEBAR ══════════ */
aside{
  width:230px;min-width:230px;
  background:#ffffff;
  border-right:1px solid #e0e3eb;
  padding:28px 18px 20px;
  display:flex;flex-direction:column;gap:16px;
  position:sticky;top:0;height:100vh;overflow-y:auto;
}
aside .logo{
  display:flex;align-items:center;gap:8px;
  font-size:16px;font-weight:800;color:#1a1a2e;
  margin-bottom:4px;
}
aside .logo span{font-size:22px}
aside h2{font-size:11px;font-weight:700;color:#999;
          text-transform:uppercase;letter-spacing:.8px}
aside label{font-size:12px;font-weight:600;color:#444;
             display:block;margin-bottom:5px}

.key-wrap{position:relative}
.key-wrap input{
  width:100%;padding:8px 32px 8px 10px;
  border:1px solid #ddd;border-radius:6px;
  font-size:12.5px;background:#fafafa;outline:none;
  transition:border-color .15s;
}
.key-wrap input:focus{border-color:#6c63ff;background:#fff}
.key-wrap button{
  position:absolute;right:8px;top:50%;
  transform:translateY(-50%);
  background:none;border:none;cursor:pointer;
  font-size:14px;color:#aaa;padding:0;line-height:1;
}
.key-wrap button:hover{color:#555}

aside hr{border:none;border-top:1px solid #eee}

.limits{
  font-size:11px;color:#aaa;line-height:1.8;
  padding:10px 12px;background:#f7f8fa;
  border-radius:6px;border:1px solid #eee;
}
.limits strong{color:#777;font-size:10px;
               text-transform:uppercase;letter-spacing:.5px;
               display:block;margin-bottom:4px}

.key-alert{
  font-size:11.5px;background:#fffde7;
  border:1px solid #f0d060;border-radius:6px;
  padding:8px 10px;color:#7a6000;line-height:1.5;
  display:none;
}
.key-alert.show{display:block}

/* ══════════ MAIN ══════════ */
main{flex:1;padding:32px 36px 40px;overflow-y:auto;max-width:1200px}

.page-title{margin-bottom:6px}
.page-title h1{font-size:26px;font-weight:800;
                display:flex;align-items:center;gap:10px}
.page-title p{font-size:13px;color:#666;margin-top:5px}

/* ══════════ GRID ══════════ */
.grid{
  display:grid;
  grid-template-columns:1fr 1fr;
  gap:28px;
  margin-top:24px;
}

.panel{
  background:#fff;border:1px solid #e0e3eb;
  border-radius:10px;padding:22px;
}
.panel-title{
  font-size:13px;font-weight:700;color:#1a1a2e;
  margin-bottom:16px;display:flex;align-items:center;gap:6px;
}
.panel-title .badge{
  font-size:10px;background:#6c63ff;color:#fff;
  padding:2px 7px;border-radius:20px;font-weight:600;
}

/* ── upload zone ── */
.drop{
  border:2px dashed #c5cae9;border-radius:8px;
  background:#fafbff;padding:22px 16px;
  text-align:center;cursor:pointer;
  position:relative;transition:all .15s;
  margin-bottom:18px;
}
.drop:hover,.drop.over{
  border-color:#6c63ff;background:#f0eeff;
}
.drop input{
  position:absolute;inset:0;opacity:0;
  cursor:pointer;width:100%;height:100%;
}
.drop-icon{font-size:28px;margin-bottom:6px}
.drop-text{font-size:13px;font-weight:600;color:#444}
.drop-sub{font-size:11px;color:#aaa;margin-top:3px}
.drop-name{
  font-size:12px;color:#6c63ff;font-weight:700;
  margin-top:8px;display:none;
}
.drop-name.show{display:block}

/* ── textarea ── */
.jd-wrap{position:relative}
textarea{
  width:100%;height:220px;
  border:1px solid #ddd;border-radius:8px;
  padding:10px 12px;font-size:13px;font-family:inherit;
  resize:vertical;background:#fafafa;color:#222;
  outline:none;line-height:1.55;
  transition:border-color .15s;
}
textarea:focus{border-color:#6c63ff;background:#fff}
.char-bar{
  display:flex;justify-content:space-between;
  margin-top:5px;font-size:11px;color:#aaa;
}
.char-bar span.warn{color:#e65100}

/* ── run button ── */
.btn-run{
  display:flex;align-items:center;justify-content:center;gap:8px;
  width:100%;margin-top:16px;padding:13px;
  background:linear-gradient(135deg,#6c63ff 0%,#48c6ef 100%);
  color:#fff;border:none;border-radius:8px;
  font-size:14px;font-weight:700;cursor:pointer;
  letter-spacing:.2px;transition:opacity .15s,transform .1s;
  box-shadow:0 2px 8px rgba(108,99,255,.25);
}
.btn-run:hover:not(:disabled){opacity:.92;transform:translateY(-1px)}
.btn-run:active:not(:disabled){transform:translateY(0)}
.btn-run:disabled{
  background:linear-gradient(135deg,#c5cae9,#b3e5fc);
  cursor:not-allowed;box-shadow:none;opacity:.8;
}

/* ── messages ── */
.msg{
  margin-top:12px;padding:10px 13px;
  border-radius:7px;font-size:12.5px;
  line-height:1.55;display:none;
}
.msg.show{display:block}
.msg.info   {background:#e8f4fd;border:1px solid #b3d9f5;color:#1565c0}
.msg.error  {background:#fff0f0;border:1px solid #ffb3b3;color:#b00020}
.msg.success{background:#f0fff4;border:1px solid #b2dfdb;color:#1b5e20}
.msg.warning{background:#fff8e1;border:1px solid #ffe082;color:#e65100}

/* ── progress ── */
.progress-wrap{margin-top:14px;display:none}
.progress-wrap.show{display:block}

.spin-row{
  display:flex;align-items:center;gap:9px;
  font-size:12.5px;color:#1565c0;margin-bottom:12px;
}
.spin{
  width:15px;height:15px;flex-shrink:0;
  border:2.5px solid #b3d9f5;border-top-color:#1565c0;
  border-radius:50%;animation:sp .7s linear infinite;
}
@keyframes sp{to{transform:rotate(360deg)}}

.steps{display:flex;flex-direction:column;gap:5px}
.step{
  display:flex;align-items:center;gap:9px;
  font-size:12px;color:#bbb;padding:4px 0;
}
.step.active{color:#1565c0;font-weight:600}
.step.done  {color:#2e7d32}
.step.fail  {color:#b00020}
.dot{
  width:7px;height:7px;border-radius:50%;
  background:#ddd;flex-shrink:0;transition:background .2s;
}
.step.active .dot{background:#1565c0}
.step.done   .dot{background:#2e7d32}
.step.fail   .dot{background:#b00020}

/* ══════════ PREVIEW PANEL ══════════ */
.preview-panel{
  background:#fff;border:1px solid #e0e3eb;
  border-radius:10px;overflow:hidden;
  display:flex;flex-direction:column;
}
.preview-header{
  padding:14px 18px;border-bottom:1px solid #eee;
  display:flex;align-items:center;justify-content:space-between;
}
.preview-header span{
  font-size:13px;font-weight:700;
  display:flex;align-items:center;gap:6px;
}
.preview-tabs{display:flex;gap:4px}
.tab{
  padding:5px 12px;border-radius:5px;font-size:11.5px;
  font-weight:600;cursor:pointer;border:1px solid #eee;
  background:#f7f8fa;color:#888;transition:all .15s;
}
.tab.active{background:#6c63ff;color:#fff;border-color:#6c63ff}

.preview-body{flex:1;min-height:500px;position:relative}
.ph{
  position:absolute;inset:0;display:flex;
  align-items:center;justify-content:center;
  padding:30px;text-align:center;color:#aaa;
  flex-direction:column;gap:10px;
}
.ph .ph-icon{font-size:40px;opacity:.4}
.ph p{font-size:13px;line-height:1.6}

iframe#pframe{
  width:100%;height:100%;min-height:500px;
  border:none;display:none;
}

.preview-footer{
  border-top:1px solid #eee;padding:12px 18px;
  display:none;align-items:center;gap:10px;
}
.preview-footer.show{display:flex}

.btn-dl{
  flex:1;display:flex;align-items:center;
  justify-content:center;gap:7px;
  padding:10px 16px;background:#1b5e20;
  color:#fff;border:none;border-radius:7px;
  font-size:13px;font-weight:700;cursor:pointer;
  text-decoration:none;transition:background .15s;
}
.btn-dl:hover{background:#2e7d32}

.ats-badge{
  display:flex;align-items:center;gap:5px;
  background:#f0fff4;border:1px solid #b2dfdb;
  color:#1b5e20;border-radius:6px;
  padding:6px 12px;font-size:11.5px;font-weight:600;
}

/* ── ATS tips panel ── */
.tips-panel{
  margin-top:16px;padding:14px 16px;
  background:#f8f9ff;border:1px solid #e0e3eb;
  border-radius:8px;display:none;
}
.tips-panel.show{display:block}
.tips-panel h4{
  font-size:12px;font-weight:700;color:#555;
  text-transform:uppercase;letter-spacing:.5px;margin-bottom:10px;
}
.tip{
  display:flex;align-items:flex-start;gap:7px;
  font-size:12px;color:#444;margin-bottom:6px;line-height:1.5;
}
.tip-icon{flex-shrink:0;font-size:13px}

/* ── responsive ── */
@media(max-width:920px){
  .grid{grid-template-columns:1fr}
  aside{width:200px;min-width:200px}
  main{padding:20px}
}
@media(max-width:620px){
  body{flex-direction:column}
  aside{
    width:100%;height:auto;position:static;
    border-right:none;border-bottom:1px solid #eee;
  }
}
</style>
</head>
<body>

<!-- ══════════════ SIDEBAR ══════════════ -->
<aside>
  <div class="logo"><span>🪄</span>ResumeAI</div>
  <hr/>

  <div>
    <h2>Configuration</h2>
  </div>

  <div>
    <label for="apiKey">Gemini API Key</label>
    <div class="key-wrap">
      <input type="password" id="apiKey"
             placeholder="AIza…" autocomplete="off"/>
      <button onclick="toggleKey()" title="Show / hide">👁</button>
    </div>
  </div>

  <div class="key-alert show" id="keyAlert">
    🔑 Enter your Gemini API key to enable optimization.
  </div>

  <hr/>

  <div class="limits">
    <strong>Limits</strong>
    Resume ≤ 15 000 chars<br>
    JD ≤ 8 000 chars<br>
    Timeout: 120 s<br>
    Upload: 10 MB max
  </div>

  <hr/>

  <div class="limits">
    <strong>How it works</strong>
    1. Upload your DOCX<br>
    2. Paste the Job Description<br>
    3. Click Tailor My Resume<br>
    4. Download ATS-ready PDF
  </div>
</aside>

<!-- ══════════════ MAIN ══════════════ -->
<main>

  <div class="page-title">
    <h1>🪄 Resume Optimizer</h1>
    <p>Upload your resume · paste the Job Description · get an ATS-optimized resume in seconds.</p>
  </div>

  <div class="grid">

    <!-- ── LEFT: inputs ── -->
    <div>

      <div class="panel">
        <div class="panel-title">
          <span>📄</span> Upload Resume
          <span class="badge">DOCX</span>
        </div>
        <div class="drop" id="drop">
          <input type="file" id="fileIn" accept=".docx"
                 onchange="onFile(this)"/>
          <div class="drop-icon">📤</div>
          <div class="drop-text">Click or drag &amp; drop your resume</div>
          <div class="drop-sub">.docx format only</div>
          <div class="drop-name" id="fname"></div>
        </div>

        <div class="panel-title" style="margin-top:4px">
          <span>📋</span> Job Description
        </div>
        <div class="jd-wrap">
          <textarea id="jd"
            placeholder="Paste the full job description here — include requirements, responsibilities, and any keywords you see…"
            oninput="onJd()"></textarea>
          <div class="char-bar">
            <span id="jd-hint" class="msg info show" style="
              margin:0;padding:5px 9px;font-size:11px;border-radius:5px">
              More detail = better ATS matching
            </span>
            <span id="cc">0 / 8 000</span>
          </div>
        </div>

        <button class="btn-run" id="run" onclick="go()" disabled>
          ✨ Tailor My Resume
        </button>

        <div class="msg info show" id="hint">
          Upload a resume, paste a job description, and enter your API key to begin.
        </div>

        <!-- progress -->
        <div class="progress-wrap" id="prog">
          <div class="spin-row">
            <div class="spin"></div>
            <span id="spinLabel">Working…</span>
          </div>
          <div class="steps">
            <div class="step" id="s1"><div class="dot"></div>📄 Extracting resume text</div>
            <div class="step" id="s2"><div class="dot"></div>🔍 Analyzing job description</div>
            <div class="step" id="s3"><div class="dot"></div>🤖 Gemini AI rewriting</div>
            <div class="step" id="s4"><div class="dot"></div>🧹 Sanitizing &amp; validating</div>
            <div class="step" id="s5"><div class="dot"></div>📄 Generating ATS PDF</div>
          </div>
        </div>

        <div class="msg" id="result"></div>
      </div>

      <!-- ATS tips -->
      <div class="tips-panel" id="tips">
        <h4>✅ ATS Optimization Applied</h4>
        <div class="tip"><span class="tip-icon">🎯</span>Keywords matched to job description</div>
        <div class="tip"><span class="tip-icon">📊</span>Achievements quantified where possible</div>
        <div class="tip"><span class="tip-icon">💪</span>Action verbs strengthened</div>
        <div class="tip"><span class="tip-icon">🔤</span>Clean formatting — no tables or images</div>
        <div class="tip"><span class="tip-icon">📋</span>Standard section headings used</div>
      </div>

    </div>

    <!-- ── RIGHT: preview ── -->
    <div>
      <div class="preview-panel">
        <div class="preview-header">
          <span>👁 Optimized Resume Preview</span>
        </div>

        <div class="preview-body">
          <div class="ph" id="ph">
            <div class="ph-icon">📄</div>
            <p>Your ATS-optimized resume will appear here.<br>
               It will closely match what the PDF looks like.</p>
          </div>
          <iframe id="pframe" title="Resume Preview"
                  sandbox="allow-same-origin"></iframe>
        </div>

        <div class="preview-footer" id="pfooter">
          <a class="btn-dl" id="dllink" href="#"
             download="ATS_Resume.pdf">
            📥 Download ATS PDF
          </a>
          <div class="ats-badge">✅ ATS Ready</div>
        </div>
      </div>
    </div>

  </div><!-- /grid -->
</main>

<!-- ══════════════ JS ══════════════ -->
<script>
"use strict";

let picked = null;

/* ── helpers ── */
const el  = id => document.getElementById(id);
const val = id => el(id).value.trim();

function toggleKey(){
  const i=el("apiKey");
  i.type=i.type==="password"?"text":"password";
}

/* ── readiness guard ── */
function guard(){
  const ok = picked && val("jd") && val("apiKey");
  el("run").disabled=!ok;

  el("keyAlert").className=val("apiKey")?"key-alert":"key-alert show";

  const hint=el("hint");
  if(!ok){
    const m=[];
    if(!picked)       m.push("upload a .docx resume");
    if(!val("jd"))    m.push("paste the job description");
    if(!val("apiKey"))m.push("enter your Gemini API key");
    hint.textContent="To continue: "+m.join(", ")+".";
    hint.className="msg info show";
  } else {
    hint.className="msg";
  }
}

function onFile(inp){
  picked=inp.files[0]||null;
  const fn=el("fname");
  if(picked){
    fn.textContent="✔ "+picked.name;
    fn.className="drop-name show";
  } else {
    fn.className="drop-name";
  }
  guard();
}

function onJd(){
  const n=el("jd").value.length;
  const cc=el("cc");
  cc.textContent=n.toLocaleString()+" / 8 000";
  cc.className=n>8000?"warn":"";
  guard();
}

el("apiKey").addEventListener("input",guard);

/* ── drag-drop ── */
const dz=el("drop");
["dragenter","dragover"].forEach(ev=>dz.addEventListener(ev,e=>{
  e.preventDefault();dz.classList.add("over");
}));
["dragleave","drop"].forEach(ev=>dz.addEventListener(ev,e=>{
  e.preventDefault();dz.classList.remove("over");
}));
dz.addEventListener("drop",e=>{
  const f=e.dataTransfer.files[0];
  if(f&&f.name.toLowerCase().endsWith(".docx")){
    picked=f;
    el("fname").textContent="✔ "+f.name;
    el("fname").className="drop-name show";
    guard();
  }
});

/* ── steps ── */
function resetSteps(){for(let i=1;i<=5;i++)el("s"+i).className="step"}
function step(n,c){el("s"+n).className="step "+c}

/* ── message ── */
function msg(id,cls,html){
  const e=el(id);
  e.className="msg "+cls+" show";
  e.innerHTML=html;
}

/* ── preview ── */
function showPreview(html,pdfUrl){
  el("ph").style.display="none";
  const fr=el("pframe");
  fr.style.display="block";

  const d=fr.contentDocument||fr.contentWindow.document;
  d.open();
  d.write(`<!DOCTYPE html><html><head><style>
    /* mirrors PDF exactly */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    *{box-sizing:border-box;margin:0;padding:0}
    body{
      font-family:Helvetica,Arial,sans-serif;
      padding:24px 28px;font-size:10.5pt;
      color:#111;line-height:1.55;
      background:#fff;
    }
    h1{
      font-size:18pt;font-weight:700;color:#000;
      border-bottom:1.5px solid #111;
      padding-bottom:5px;margin-bottom:10px;
    }
    h2{
      font-size:12pt;font-weight:700;color:#111;
      border-bottom:0.75px solid #bbb;
      padding-bottom:3px;margin-top:14px;margin-bottom:7px;
      text-transform:uppercase;letter-spacing:.4px;
    }
    h3{
      font-size:10.5pt;font-weight:700;color:#222;
      margin-top:8px;margin-bottom:2px;
    }
    p{margin-bottom:5px}
    ul{padding-left:18px;margin:4px 0 8px}
    li{margin-bottom:3px}
    strong,b{font-weight:700;color:#000}
    em,i{font-style:italic;color:#333}
    .section{margin-bottom:12px}
  </style></head><body>${html}</body></html>`);
  d.close();

  const pf=el("pfooter");
  if(pdfUrl){
    el("dllink").href=pdfUrl;
    pf.className="preview-footer show";
  } else {
    pf.className="preview-footer";
  }

  el("tips").className="tips-panel show";
}

/* ══════════ MAIN PIPELINE ══════════ */
async function go(){
  if(!picked)return;

  /* reset */
  el("result").className="msg";
  el("prog").className="progress-wrap show";
  el("tips").className="tips-panel";
  resetSteps();
  el("run").disabled=true;
  el("ph").style.display="flex";
  el("pframe").style.display="none";
  el("pfooter").className="preview-footer";

  const fd=new FormData();
  fd.append("resume",  picked);
  fd.append("jd_text", el("jd").value);
  fd.append("api_key", val("apiKey"));

  /* step animations while waiting */
  step(1,"active");
  el("spinLabel").textContent="📄 Extracting resume…";

  setTimeout(()=>{
    step(1,"done");step(2,"active");
    el("spinLabel").textContent="🔍 Analyzing job description…";
  },600);
  setTimeout(()=>{
    step(2,"done");step(3,"active");
    el("spinLabel").textContent="🤖 Gemini AI is rewriting your resume…";
  },1400);

  /* abort controller — 150 s browser-side timeout */
  const ctrl=new AbortController();
  const tid=setTimeout(()=>ctrl.abort(),150_000);

  let data;
  try{
    const resp=await fetch("/optimize",{
      method:"POST",body:fd,signal:ctrl.signal
    });
    clearTimeout(tid);

    const raw=await resp.text();
    try{ data=JSON.parse(raw); }
    catch(_){
      throw new Error(
        "Server returned non-JSON (HTTP "+resp.status+").\n"+raw.slice(0,400)
      );
    }
  }catch(err){
    clearTimeout(tid);
    el("prog").className="progress-wrap";
    for(let i=1;i<=5;i++){
      if(el("s"+i).classList.contains("active")){step(i,"fail");break;}
    }
    const txt=err.name==="AbortError"
      ? "⏱ Timed out (>150 s). Try shorter inputs or retry."
      : "❌ "+err.message;
    msg("result","error",txt);
    el("run").disabled=false;
    return;
  }

  /* update steps from server */
  const reached=data.steps_reached||1;
  for(let i=1;i<reached;i++) step(i,"done");
  step(reached, data.ok?"done":"fail");
  if(data.ok) for(let i=reached+1;i<=5;i++) step(i,"done");

  el("prog").className="progress-wrap";

  if(!data.ok){
    msg("result","error","❌ "+( data.error||"Unknown error." ));
    el("run").disabled=false;
    return;
  }

  showPreview(data.html, data.pdf_url||null);

  if(data.pdf_url){
    msg("result","success",
      "✅ ATS-optimized resume ready! Preview shown above — download your PDF below.");
  } else {
    msg("result","warning",
      "✅ Resume generated. PDF export failed: "+(data.pdf_warning||"")+
      "<br>You can still copy text from the preview.");
  }
  el("run").disabled=false;
}

guard();
</script>
</body>
</html>"""


# ═══════════════════════════════════════════════════════════════
# PDF UTILITIES
# ═══════════════════════════════════════════════════════════════

class HRule(Flowable):
    """Horizontal rule — mirrors CSS border-bottom on headings."""
    def __init__(self, thickness=0.75, color=colors.HexColor("#222"),
                 space_before=0, space_after=6, width_frac=1.0):
        super().__init__()
        self.thickness   = thickness
        self.line_color  = color
        self.space_before= space_before
        self.space_after = space_after
        self.width_frac  = width_frac

    def wrap(self, aw, _ah):
        self._aw = aw
        return aw, self.thickness + self.space_before + self.space_after

    def draw(self):
        c = self.canv
        c.setStrokeColor(self.line_color)
        c.setLineWidth(self.thickness)
        y = self.space_after
        c.line(0, y, self._aw * self.width_frac, y)


# ═══════════════════════════════════════════════════════════════
# HTML PARSER  (inline bold/italic preserved for ReportLab)
# ═══════════════════════════════════════════════════════════════

class _Parser(HTMLParser):
    """
    Converts our HTML subset to (tag, rl_markup) tuples.
    Inline <strong>/<em> → <b>/<i> for ReportLab Paragraph XML.
    """
    BLOCKS  = {"h1","h2","h3","p","li","ul"}
    INLINES = {"strong":"b","em":"i","b":"b","i":"i"}

    def __init__(self):
        super().__init__()
        self._stack: list[str] = []
        self._buf   = ""
        self.nodes: list[tuple[str,str]] = []

    def _flush(self, tag: str):
        t = self._buf.strip()
        self._buf = ""
        if t:
            self.nodes.append((tag, t))

    def _cur_block(self):
        for t in reversed(self._stack):
            if t in self.BLOCKS:
                return t
        return "p"

    def handle_starttag(self, tag, _):
        tag = tag.lower()
        if tag in self.BLOCKS:
            if self._buf.strip():
                self._flush(self._cur_block())
            self._stack.append(tag)
        elif tag in self.INLINES:
            self._buf += f"<{self.INLINES[tag]}>"
            self._stack.append(tag)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in self.BLOCKS:
            self._flush(tag)
        elif tag in self.INLINES:
            self._buf += f"</{self.INLINES[tag]}>"
        while self._stack and self._stack[-1] != tag:
            self._stack.pop()
        if self._stack:
            self._stack.pop()

    def handle_data(self, data):
        self._buf += data

    def handle_entityref(self, name):
        self._buf += {
            "amp":"&amp;","lt":"&lt;","gt":"&gt;",
            "nbsp":" ","quot":'"',"apos":"'"
        }.get(name, "")

    def handle_charref(self, name):
        try:
            self._buf += chr(
                int(name[1:], 16) if name.startswith("x") else int(name)
            )
        except Exception:
            pass

    def run(self, html: str):
        self.feed(html)
        if self._buf.strip():
            self._flush(self._cur_block())
        return self.nodes


# ═══════════════════════════════════════════════════════════════
# DOCX EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_docx(raw: bytes):
    try:
        document = docx.Document(io.BytesIO(raw))
        lines: list[str] = []
        for p in document.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in lines:
                        lines.append(t)
        return "\n".join(lines).strip(), None
    except Exception as e:
        return None, f"Cannot read DOCX: {e}"


# ═══════════════════════════════════════════════════════════════
# VALIDATION
# ═══════════════════════════════════════════════════════════════

def check_resume(text: str):
    if not text:
        return False, "No text extracted — file may be empty or image-based."
    if len(text) < RESUME_MIN:
        return False, f"Resume too short ({len(text)} chars). Check the DOCX."
    if len(text) > RESUME_MAX:
        return False, f"Resume too long ({len(text):,} chars). Max {RESUME_MAX:,}."
    return True, ""

def check_jd(text: str):
    s = text.strip()
    if not s:
        return False, "Job description is empty."
    if len(s) > JD_MAX:
        return False, f"JD too long ({len(s):,} chars). Max {JD_MAX:,}."
    return True, ""


# ═══════════════════════════════════════════════════════════════
# GEMINI API
# ═══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """
You are a world-class ATS resume writer and career coach.

Your task: Rewrite the provided resume so it is perfectly tailored to the Job Description (JD).

STRICT OUTPUT RULES:
1. Output ONLY valid HTML. No markdown. No code fences. No explanations.
2. Allowed tags ONLY: <h1> <h2> <h3> <p> <ul> <li> <strong> <em>
3. No <html>, <head>, <body>, <style>, <script> tags.
4. No inline CSS, no attributes on any tag.
5. No tables, no columns, no images.

CONTENT RULES:
1. Start with <h1>Full Name</h1> then contact info in <p>.
2. Use <h2> for section headings: SUMMARY, EXPERIENCE, EDUCATION, SKILLS, CERTIFICATIONS.
3. Under EXPERIENCE: <h3> for each job title + company + dates.
4. Use <ul><li> for bullet points under each role and skills.
5. Mirror keywords, phrases, and terminology from the JD exactly.
6. Strengthen all bullet points with action verbs (Led, Built, Delivered, etc.).
7. Quantify achievements wherever possible (%, $, numbers).
8. Add a SKILLS section with keywords extracted directly from the JD.
9. Write a 3-sentence SUMMARY that matches the job title and requirements.
10. NEVER invent facts, employers, degrees, or dates not in the original resume.
11. Remove irrelevant experience that does not relate to the JD.
12. Keep formatting clean — ATS systems cannot read fancy layouts.
""".strip()


def gemini_err(resp) -> str:
    try:
        return resp.json().get("error",{}).get("message","Unknown.")
    except Exception:
        return resp.text[:300] if resp.text else "(empty)"


def call_gemini(key: str, resume: str, jd: str):
    url = (
        "https://generativelanguage.googleapis.com/v1beta/"
        f"models/gemini-2.5-flash:generateContent?key={key}"
    )
    body = {
        "contents": [{"role":"user","parts":[{
            "text": (
                f"ORIGINAL RESUME:\n\n{resume}\n\n"
                f"{'─'*60}\n\n"
                f"JOB DESCRIPTION:\n\n{jd}"
            )
        }]}],
        "systemInstruction": {"parts":[{"text": SYSTEM_PROMPT}]},
        "generationConfig": {
            "temperature":      0.3,
            "maxOutputTokens":  8192,
            "topP":             0.85,
        },
    }
    try:
        r = req.post(
            url,
            headers={"Content-Type":"application/json"},
            data=json.dumps(body),
            timeout=API_TIMEOUT,
        )
    except req.exceptions.ConnectionError:
        return None, "Cannot reach Gemini API. Check internet connection."
    except req.exceptions.Timeout:
        return None, f"Gemini timed out after {API_TIMEOUT}s. Try again."
    except req.exceptions.RequestException as e:
        return None, f"Request error: {e}"

    if not r.ok:
        m = gemini_err(r)
        mapping = {
            401: "Invalid API key. Check Settings.",
            403: "API key lacks permission for this model.",
            429: "Rate limit reached. Wait a minute and retry.",
            400: f"Bad request: {m}",
        }
        return None, mapping.get(r.status_code, f"HTTP {r.status_code}: {m}")

    try:
        return r.json(), None
    except Exception:
        return None, "Gemini returned invalid JSON."


# ═══════════════════════════════════════════════════════════════
# RESPONSE PARSING
# ═══════════════════════════════════════════════════════════════

def parse_response(data: dict):
    cands = data.get("candidates", [])
    if not cands:
        return None, "No candidates in Gemini response."

    first   = cands[0]
    reason  = first.get("finishReason","UNKNOWN")

    if reason in ("SAFETY","RECITATION"):
        return None, f"Gemini blocked output (reason: {reason}). Try rephrasing."

    content = first.get("content", {})
    parts   = content.get("parts", [])
    if not parts:
        return None, f"No content from Gemini (finishReason: {reason})."

    text = parts[0].get("text","").strip()
    if not text:
        return None, "Gemini returned empty text."

    # detect refusal
    refusals = ["i cannot","i'm unable","i am unable",
                "i can't","i won't","as an ai"]
    if any(r in text.lower() for r in refusals) and len(text) < 500:
        return None, "Gemini refused. Try adjusting the resume or JD."

    return text, None


# ═══════════════════════════════════════════════════════════════
# HTML CLEANING
# ═══════════════════════════════════════════════════════════════

def clean_html(raw: str):
    # strip markdown fences
    text = re.sub(r"^\s*```[a-zA-Z]*\s*\n?", "", raw, flags=re.MULTILINE)
    text = re.sub(r"\n?\s*```\s*$", "", text, flags=re.MULTILINE)

    # strip any preamble before first HTML tag
    m = re.search(r"<[a-zA-Z]", text)
    if m and m.start() > 0:
        text = text[m.start():]

    # bleach allowlist
    text = bleach.clean(
        text,
        tags=ALLOWED_TAGS,
        attributes=ALLOWED_ATTRS,
        strip=True,
        strip_comments=True,
    )
    text = text.strip()

    if not text:
        return None, "Empty HTML after cleaning."
    if not re.search(r"<(h[123]|p|li)\b", text, re.I):
        return None, "AI output has no resume structure. Try again."
    if len(re.sub(r"<[^>]+>","",text).strip()) < 50:
        return None, "AI output has too little text. Try again."

    return text, None


# ═══════════════════════════════════════════════════════════════
# PDF GENERATION  — ReportLab
# ═══════════════════════════════════════════════════════════════

def make_pdf(html: str):
    try:
        def S(name, **kw):
            return ParagraphStyle(name, **kw)

        # ── styles that mirror the iframe CSS exactly ──
        styles = {
            "h1": S("H1",
                fontName="Helvetica-Bold", fontSize=18, leading=22,
                textColor=colors.HexColor("#000000"),
                spaceBefore=0, spaceAfter=2,
            ),
            "h2": S("H2",
                fontName="Helvetica-Bold", fontSize=12, leading=16,
                textColor=colors.HexColor("#111111"),
                spaceBefore=14, spaceAfter=2,
            ),
            "h3": S("H3",
                fontName="Helvetica-Bold", fontSize=10.5, leading=14,
                textColor=colors.HexColor("#222222"),
                spaceBefore=8, spaceAfter=2,
            ),
            "p": S("P",
                fontName="Helvetica", fontSize=10.5, leading=14,
                textColor=colors.HexColor("#111111"),
                spaceAfter=5, alignment=TA_LEFT,
            ),
            "li": S("LI",
                fontName="Helvetica", fontSize=10.5, leading=14,
                textColor=colors.HexColor("#111111"),
                spaceAfter=3, leftIndent=12,
            ),
        }

        nodes  = _Parser().run(html)
        story  = []
        li_buf : list[str] = []

        def flush_list():
            if not li_buf:
                return
            items = [
                ListItem(
                    Paragraph(t, styles["li"]),
                    leftIndent=14,
                    bulletColor=colors.HexColor("#444"),
                )
                for t in li_buf
            ]
            story.append(
                ListFlowable(
                    items,
                    bulletType="bullet",
                    leftIndent=14,
                    spaceBefore=3,
                    spaceAfter=8,
                )
            )
            li_buf.clear()

        for tag, markup in nodes:
            if tag == "li":
                li_buf.append(markup)
                continue

            flush_list()

            story.append(Paragraph(markup, styles.get(tag, styles["p"])))

            # draw horizontal rules after headings — matches CSS border-bottom
            if tag == "h1":
                story.append(
                    HRule(thickness=1.5,
                          color=colors.HexColor("#111111"),
                          space_before=2, space_after=8)
                )
            elif tag == "h2":
                story.append(
                    HRule(thickness=0.75,
                          color=colors.HexColor("#bbbbbb"),
                          space_before=2, space_after=6)
                )

        flush_list()

        if not story:
            return None, "No content to render in PDF."

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=1.8*cm, rightMargin=1.8*cm,
            topMargin=1.5*cm,  bottomMargin=1.5*cm,
        )
        doc.build(story)

        pdf = buf.getvalue()
        if not pdf:
            return None, "PDF file was empty after rendering."
        return pdf, None

    except Exception as e:
        return None, f"PDF error: {e}\n{traceback.format_exc()}"


# ═══════════════════════════════════════════════════════════════
# PDF STORE  (in-memory)
# ═══════════════════════════════════════════════════════════════

_store : dict[str, bytes] = {}
_lock  = threading.Lock()


def save_pdf(data: bytes) -> str:
    key = hashlib.sha256(
        (str(time.time()) + str(id(data))).encode()
    ).hexdigest()[:20]
    with _lock:
        _store[key] = data
    return key


# ═══════════════════════════════════════════════════════════════
# FLASK ROUTES
# ═══════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template_string(PAGE)


@app.route("/optimize", methods=["POST","OPTIONS"])
def optimize():
    if request.method == "OPTIONS":
        return "", 204

    def fail(msg: str, step: int):
        return jsonify({"ok": False, "error": msg, "steps_reached": step})

    try:
        # ── step 1: extract & validate ──────────────────────────
        f = request.files.get("resume")
        if not f:
            return fail("No file uploaded. Choose a .docx file.", 1)

        raw = f.read()
        if not raw:
            return fail("Uploaded file is empty.", 1)

        text, err = extract_docx(raw)
        if err:
            return fail(err, 1)

        ok, msg = check_resume(text)
        if not ok:
            return fail(msg, 1)

        jd = request.form.get("jd_text","").strip()
        ok, msg = check_jd(jd)
        if not ok:
            return fail(msg, 1)

        api_key = request.form.get("api_key","").strip()
        if not api_key:
            return fail("API key missing. Enter it in Settings.", 1)

        # ── step 2: call Gemini ──────────────────────────────────
        resp, err = call_gemini(api_key, text, jd)
        if err:
            return fail(err, 2)

        # ── step 3: parse response ───────────────────────────────
        raw_html, err = parse_response(resp)
        if err:
            return fail(err, 3)

        # ── step 4: clean HTML ───────────────────────────────────
        html, err = clean_html(raw_html)
        if err:
            return fail(err, 4)

        # ── step 5: generate PDF ─────────────────────────────────
        pdf, err = make_pdf(html)
        if err:
            return jsonify({
                "ok":          True,
                "html":        html,
                "pdf_url":     None,
                "pdf_warning": err,
                "steps_reached": 5,
            })

        pid = save_pdf(pdf)
        return jsonify({
            "ok":          True,
            "html":        html,
            "pdf_url":     f"/download/{pid}",
            "steps_reached": 5,
        })

    except Exception as e:
        tb = traceback.format_exc()
        app.logger.error("Unhandled /optimize error:\n%s", tb)
        return fail(f"Unexpected server error: {e}", 1)


@app.route("/download/<pid>")
def download(pid: str):
    with _lock:
        pdf = _store.get(pid)
    if not pdf:
        return "PDF not found or expired.", 404
    return send_file(
        io.BytesIO(pdf),
        mimetype="application/pdf",
        as_attachment=True,
        download_name="ATS_Optimized_Resume.pdf",
    )


# ═══════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    app.run(debug=True, port=5000, threaded=True)